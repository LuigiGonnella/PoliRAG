"""PDF document loader."""
from docx import Document
from pathlib import Path
from typing import List
import os
from cleaners.text_cleaner import clean_text
from chunking.chunk_manager import line_text_splitter

CHUNK_SIZE = 500


def get_docx_file(path):
    filename = os.path.split(path)[-1]
    doc = Document(Path(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    
    text ="\n".join(parts)
    return text, filename
    

def chunk_docx(path):
    try:
        text, filename = get_docx_file(path)
        chunks = line_text_splitter(clean_text(text), CHUNK_SIZE)

    except Exception as e:
        print(f"Error encountered: {str(e)}")

    print(f"DOCUMENT {filename} chunking completed successfully!")

    return chunks