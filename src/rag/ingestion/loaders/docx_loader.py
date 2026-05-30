"""DOCX document loader."""
import os
from pathlib import Path
from docx import Document
from src.rag.ingestion.cleaners.text_cleaner import clean_text
from chunking.chunk_manager import line_text_splitter

CHUNK_SIZE = 500

def get_docx_file(path):
    doc = Document(Path(path))
    parts = []
    
    # Extract structural text paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text.strip())
            
    # Convert tabular structures cleanly to markdown cells instead of flat noisy words
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            if any(row_cells):  # Avoid empty rows
                parts.append("| " + " | ".join(row_cells) + " |")
    
    return "\n".join(parts)

def chunk_docx(path):
    filename = os.path.basename(path)
    chunks_with_metadata = []
    
    try:
        text = get_docx_file(path)
        cleaned_text = clean_text(text)
        
        # Execute text splitting pass
        raw_chunks = line_text_splitter(cleaned_text, CHUNK_SIZE)
        
        # Structure into standard dictionary list format
        for idx, chunk_text in enumerate(raw_chunks):
            chunks_with_metadata.append({
                "text": chunk_text,
                "index": idx + 1,  # Sequential chunk index since DOCX lacks hard pages
                "source": path
            })
            
    except Exception as e:
        print(f"Error encountered processing DOCX file {filename}: {str(e)}")
        
    print(f"DOCUMENT {filename} chunking completed successfully! Created {len(chunks_with_metadata)} chunks.")
    return chunks_with_metadata